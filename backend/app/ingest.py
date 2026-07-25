"""Ingestion pipeline: parse -> chunk -> embed -> store (with ACL tags)."""
import io

from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session

from app.config import settings
from app.embeddings import embed_texts
from app.models import Chunk, Document


class UnsupportedFileType(Exception):
    pass


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    """Extract raw text from an uploaded file (PDF, DOCX, or plain text)."""
    name = filename.lower()

    if name.endswith(".pdf") or content_type == "application/pdf":
        import fitz  # PyMuPDF

        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)

    if name.endswith(".docx") or "wordprocessingml" in content_type:
        import docx

        document = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)

    if name.endswith((".txt", ".md")) or content_type.startswith("text/"):
        return data.decode("utf-8", errors="replace")

    raise UnsupportedFileType(
        f"Unsupported file type: {filename} ({content_type}). "
        "Supported: .pdf, .docx, .txt, .md"
    )


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks. `chunk_size` is in characters here;
    ~512 chars is a reasonable default that keeps chunks small and focused."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return [c.strip() for c in splitter.split_text(text) if c.strip()]


def ingest_document(
    db: Session,
    *,
    filename: str,
    content_type: str,
    data: bytes,
    allowed_roles: list[str],
    uploaded_by: str,
) -> Document:
    """Run the full pipeline and persist the document + its chunks."""
    text = extract_text(filename, content_type, data)
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("No extractable text found in the document.")

    vectors = embed_texts(chunks)

    document = Document(
        filename=filename,
        content_type=content_type,
        uploaded_by=uploaded_by,
        allowed_roles=allowed_roles,
        num_chunks=len(chunks),
    )
    db.add(document)
    db.flush()  # assigns document.id without committing yet

    for i, (chunk, vector) in enumerate(zip(chunks, vectors)):
        db.add(
            Chunk(
                document_id=document.id,
                chunk_index=i,
                content=chunk,
                allowed_roles=allowed_roles,  # denormalized for fast ACL filtering
                embedding=vector,
            )
        )

    db.commit()
    db.refresh(document)
    return document
